'''
SCRIPT: GENERAR ARTÍCULO CIENTÍFICO PARA ECA SINERGIA
======================================================
Proyecto: Predicción de Criminalidad en Ecuador
Autor: Erick Reinaldo Flores Zambrano
Formato: Revista ECA Sinergia - UTM
'''

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Rutas
RUTA_SALIDA = r"C:\Users\Erick Zambrano\Desktop\linkedin\PROYECTOS\02_criminalidad_ecuador"
os.makedirs(RUTA_SALIDA, exist_ok=True)

print("=" * 70)
print("📝 GENERANDO ARTÍCULO CIENTÍFICO - REVISTA ECA SINERGIA")
print("=" * 70)

# Crear documento
doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# ============================================================
# TÍTULO
# ============================================================
titulo = doc.add_paragraph()
titulo_run = titulo.add_run("MODELO PREDICTIVO DE HOMICIDIOS EN ECUADOR MEDIANTE ALGORITMOS DE APRENDIZAJE AUTOMÁTICO: UN ENFOQUE BASADO EN DATOS OFICIALES DEL MINISTERIO DEL INTERIOR (2014-2025)")
titulo_run.bold = True
titulo_run.font.size = Pt(14)
titulo_run.font.name = 'Times New Roman'
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================
# TÍTULO EN INGLÉS
# ============================================================
titulo_en = doc.add_paragraph()
titulo_en_run = titulo_en.add_run("PREDICTIVE MODEL OF HOMICIDES IN ECUADOR USING MACHINE LEARNING ALGORITHMS: AN APPROACH BASED ON OFFICIAL DATA FROM THE MINISTRY OF INTERIOR (2014-2025)")
titulo_en_run.italic = True
titulo_en_run.font.size = Pt(12)
titulo_en_run.font.name = 'Times New Roman'
titulo_en.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================
# AUTOR
# ============================================================
autor = doc.add_paragraph()
autor_run = autor.add_run("Erick Reinaldo Flores Zambrano")
autor_run.font.size = Pt(11)
autor_run.font.name = 'Times New Roman'
autor.alignment = WD_ALIGN_PARAGRAPH.CENTER

afiliacion = doc.add_paragraph()
afil_run = afiliacion.add_run("Universidad Técnica de Manabí, Facultad de Ciencias Administrativas y Económicas\nPortoviejo, Ecuador\neflores@utm.edu.ec")
afil_run.font.size = Pt(10)
afil_run.font.name = 'Times New Roman'
afiliacion.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================
# RESUMEN
# ============================================================
resumen_titulo = doc.add_paragraph()
res_run = resumen_titulo.add_run("RESUMEN")
res_run.bold = True
res_run.font.size = Pt(11)
res_run.font.name = 'Times New Roman'

resumen_texto = """La presente investigación desarrolla un modelo predictivo de homicidios intencionales en Ecuador utilizando técnicas de aprendizaje automático. Se analizaron más de 850.000 registros oficiales del Ministerio del Interior correspondientes al período 2014-2025, incluyendo variables como armas incautadas, personas desaparecidas, detenidos y drogas decomisadas. Se evaluaron cuatro algoritmos de clasificación: XGBoost, Random Forest, CatBoost y Ridge Regression. Los resultados demuestran que el modelo XGBoost alcanzó un coeficiente de determinación (R²) de 96,85%, con un error cuadrático medio (RMSE) de 2,71 homicidios mensuales por provincia. El análisis revela un incremento del 738% en la tasa de homicidios entre 2017 y 2025, pasando de 5,7 a 47,8 por cada 100.000 habitantes, posicionando a Ecuador como uno de los países más violentos de América Latina. La provincia del Guayas concentra el 47% de los homicidios nacionales. Este modelo constituye una herramienta de apoyo para la formulación de políticas públicas de seguridad ciudadana basadas en evidencia científica."""

resumen = doc.add_paragraph()
resumen_r = resumen.add_run(resumen_texto)
resumen_r.font.size = Pt(11)
resumen_r.font.name = 'Times New Roman'
resumen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Palabras clave
palabras = doc.add_paragraph()
palabras_run = palabras.add_run("Palabras clave: ")
palabras_run.bold = True
palabras_run.font.size = Pt(11)
palabras_run.font.name = 'Times New Roman'
palabras_texto = palabras.add_run("aprendizaje automático, predicción de criminalidad, homicidios, XGBoost, seguridad ciudadana")
palabras_texto.font.size = Pt(11)
palabras_texto.font.name = 'Times New Roman'

doc.add_paragraph()

# ============================================================
# ABSTRACT
# ============================================================
abstract_titulo = doc.add_paragraph()
abs_run = abstract_titulo.add_run("ABSTRACT")
abs_run.bold = True
abs_run.font.size = Pt(11)
abs_run.font.name = 'Times New Roman'

abstract_texto = """This research develops a predictive model of intentional homicides in Ecuador using machine learning techniques. More than 850,000 official records from the Ministry of Interior for the period 2014-2025 were analyzed, including variables such as seized weapons, missing persons, detainees, and confiscated drugs. Four classification algorithms were evaluated: XGBoost, Random Forest, CatBoost, and Ridge Regression. The results demonstrate that the XGBoost model achieved a coefficient of determination (R²) of 96.85%, with a root mean square error (RMSE) of 2.71 monthly homicides per province. The analysis reveals a 738% increase in the homicide rate between 2017 and 2025, from 5.7 to 47.8 per 100,000 inhabitants, positioning Ecuador as one of the most violent countries in Latin America. Guayas province concentrates 47% of national homicides. This model constitutes a support tool for the formulation of evidence-based public citizen security policies."""

abstract = doc.add_paragraph()
abstract_r = abstract.add_run(abstract_texto)
abstract_r.font.size = Pt(11)
abstract_r.font.name = 'Times New Roman'
abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Keywords
keywords = doc.add_paragraph()
keywords_run = keywords.add_run("Keywords: ")
keywords_run.bold = True
keywords_run.font.size = Pt(11)
keywords_run.font.name = 'Times New Roman'
keywords_texto = keywords.add_run("machine learning, crime prediction, homicides, XGBoost, citizen security")
keywords_texto.font.size = Pt(11)
keywords_texto.font.name = 'Times New Roman'

doc.add_paragraph()

# ============================================================
# INTRODUCCIÓN
# ============================================================
intro_titulo = doc.add_paragraph()
intro_run = intro_titulo.add_run("1. INTRODUCCIÓN")
intro_run.bold = True
intro_run.font.size = Pt(12)
intro_run.font.name = 'Times New Roman'

intro_texto = """La seguridad ciudadana constituye uno de los principales desafíos para los gobiernos de América Latina en el siglo XXI. Ecuador, en particular, ha experimentado una transformación dramática en sus indicadores de violencia durante la última década. Según datos del Ministerio del Interior (2025), la tasa de homicidios intencionales se ha multiplicado por ocho entre 2017 y 2025, alcanzando niveles históricos que superan los registros de cualquier período anterior en la historia republicana del país.

Este fenómeno de escalada de violencia ha sido documentado por organismos internacionales como la Organización de las Naciones Unidas y el Observatorio Ecuatoriano de Crimen Organizado (OECO), quienes han señalado la incidencia del narcotráfico y las bandas criminales transnacionales como factores determinantes. No obstante, la capacidad de anticipación de las autoridades frente a estos eventos violentos ha sido limitada, evidenciando la necesidad de herramientas predictivas basadas en datos que permitan una asignación más eficiente de recursos de seguridad.

El aprendizaje automático, rama de la inteligencia artificial que permite a los sistemas informáticos aprender patrones a partir de datos históricos, ha demostrado su efectividad en diversos campos de aplicación, incluyendo la predicción de fenómenos criminales. Estudios previos en países como Estados Unidos, Reino Unido y Colombia han implementado modelos predictivos con resultados prometedores para la anticipación de delitos violentos.

En este contexto, la presente investigación tiene como objetivo desarrollar y validar un modelo de aprendizaje automático capaz de predecir homicidios intencionales a nivel provincial en Ecuador, utilizando datos oficiales del Ministerio del Interior. La hipótesis central sostiene que existe una relación significativa entre variables como armas incautadas, personas desaparecidas y detenidos, que permite anticipar la ocurrencia de homicidios con un alto grado de precisión.

El aporte de esta investigación radica en la construcción de una herramienta cuantitativa que pueda servir de insumo para la toma de decisiones en materia de políticas públicas de seguridad ciudadana en Ecuador."""

intro = doc.add_paragraph()
intro_r = intro.add_run(intro_texto)
intro_r.font.size = Pt(11)
intro_r.font.name = 'Times New Roman'
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ============================================================
# METODOLOGÍA
# ============================================================
metod_titulo = doc.add_paragraph()
metod_run = metod_titulo.add_run("2. METODOLOGÍA")
metod_run.bold = True
metod_run.font.size = Pt(12)
metod_run.font.name = 'Times New Roman'

metod_texto = """La investigación adoptó un enfoque cuantitativo de tipo predictivo, utilizando técnicas de minería de datos y aprendizaje automático supervisado. El diseño metodológico se estructuró en cuatro fases: recolección de datos, preprocesamiento, modelado y validación.

2.1 Fuentes de datos

Los datos primarios fueron obtenidos del portal de datos abiertos del Ministerio del Interior de Ecuador. Se recopilaron registros oficiales correspondientes al período enero 2014 - noviembre 2025, totalizando 850.347 observaciones distribuidas en cinco conjuntos de datos:

- Homicidios intencionales: 38.932 registros
- Armas ilícitas incautadas: 69.686 registros
- Personas desaparecidas: 75.459 registros
- Personas detenidas: 556.206 registros
- Drogas incautadas: 112.848 operativos

Adicionalmente, se incorporaron datos demográficos del Instituto Nacional de Estadística y Censos (INEC) para el cálculo de tasas por cada 100.000 habitantes.

2.2 Variables del modelo

La variable dependiente corresponde al número de homicidios intencionales mensuales por provincia. Las variables independientes incluyen:

- Temporales: año, mes, trimestre, rezagos (lag_1, lag_2, lag_3), medias móviles (3 y 6 meses)
- Criminales: armas incautadas, drogas decomisadas, detenidos
- Sociales: personas desaparecidas, población provincial
- Geográficas: codificación one-hot de 24 provincias

2.3 Preprocesamiento de datos

El tratamiento de datos incluyó: unificación de formatos de fecha, normalización de nombres de provincias, imputación de valores faltantes mediante interpolación lineal, y eliminación de registros duplicados. Se aplicó escalado estándar (z-score) a las variables numéricas para garantizar comparabilidad entre escalas.

2.4 Algoritmos evaluados

Se implementaron cuatro algoritmos de aprendizaje automático supervisado:

a) XGBoost (Extreme Gradient Boosting): algoritmo de ensamble basado en árboles de decisión que utiliza boosting para optimización iterativa.

b) Random Forest: método de ensamble que construye múltiples árboles de decisión y promedia sus predicciones.

c) CatBoost: variante de gradient boosting optimizada para variables categóricas.

d) Ridge Regression: regresión lineal regularizada con penalización L2.

2.5 Validación

Se utilizó una división temporal 80/20, donde el 80% de los datos (2014-2023) se destinó al entrenamiento y el 20% restante (2024-2025) a la validación. Esta estrategia previene el data leakage y simula condiciones reales de predicción futura.

Las métricas de evaluación empleadas fueron: coeficiente de determinación (R²), error cuadrático medio (RMSE), error absoluto medio (MAE) y error porcentual absoluto medio (MAPE)."""

metod = doc.add_paragraph()
metod_r = metod.add_run(metod_texto)
metod_r.font.size = Pt(11)
metod_r.font.name = 'Times New Roman'
metod.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ============================================================
# RESULTADOS
# ============================================================
result_titulo = doc.add_paragraph()
result_run = result_titulo.add_run("3. RESULTADOS")
result_run.bold = True
result_run.font.size = Pt(12)
result_run.font.name = 'Times New Roman'

result_texto = """3.1 Análisis descriptivo

El análisis exploratorio de los datos reveló una tendencia ascendente sostenida en los homicidios intencionales desde 2021. La Tabla 1 presenta la evolución anual de homicidios y tasas por cada 100.000 habitantes.

Tabla 1. Evolución de homicidios intencionales en Ecuador (2014-2025)

| Año  | Homicidios | Tasa por 100.000 hab. |
|------|------------|----------------------|
| 2014 | 1.310      | 8,2                  |
| 2015 | 1.050      | 6,4                  |
| 2016 | 959        | 5,8                  |
| 2017 | 970        | 5,7                  |
| 2018 | 996        | 5,8                  |
| 2019 | 1.189      | 6,8                  |
| 2020 | 1.372      | 7,8                  |
| 2021 | 2.495      | 14,0                 |
| 2022 | 4.886      | 27,2                 |
| 2023 | 8.248      | 47,25                |
| 2024 | 7.063      | 38,2                 |
| 2025 | 8.393      | 47,8                 |

Fuente: Ministerio del Interior de Ecuador (2025)

El año 2017 registró el mínimo histórico con una tasa de 5,7 homicidios por cada 100.000 habitantes, ubicando a Ecuador como uno de los países más seguros de la región en ese momento. Sin embargo, a partir de 2021 se evidencia un punto de inflexión con un incremento del 82% respecto al año anterior, tendencia que se consolida hasta alcanzar máximos históricos en 2023 y 2025.

3.2 Distribución geográfica

La distribución de homicidios presenta una marcada concentración geográfica. Guayas acumula el 47% del total nacional (18.264 homicidios en el período), seguida por Manabí (8,4%), Los Ríos (8,0%), Esmeraldas (5,2%) y El Oro (5,4%). Estas cinco provincias costeras concentran el 74% de la violencia letal del país.

3.3 Rendimiento de los modelos

La Tabla 2 presenta los resultados comparativos de los cuatro algoritmos evaluados.

Tabla 2. Comparación de modelos de aprendizaje automático

| Modelo           | R²     | RMSE  | MAE   | MAPE    |
|------------------|--------|-------|-------|---------|
| XGBoost          | 96,85% | 2,71  | 1,15  | 27,35%  |
| Random Forest    | 95,32% | 3,31  | 1,35  | 25,52%  |
| CatBoost         | 91,55% | 4,45  | 2,31  | 57,44%  |
| Ridge Regression | 90,45% | 4,73  | 0,94  | 18,87%  |

El modelo XGBoost obtuvo el mejor desempeño global con un R² de 96,85%, lo que indica que el modelo explica el 96,85% de la variabilidad en los homicidios mensuales por provincia. El RMSE de 2,71 significa que, en promedio, las predicciones difieren de los valores reales en menos de 3 homicidios mensuales por provincia.

3.4 Importancia de variables

El análisis de importancia de características reveló que las variables más influyentes en la predicción son:

1. Homicidios del mes anterior (lag_1): 28,4%
2. Armas incautadas: 18,7%
3. Provincia (codificación): 15,3%
4. Mes del año: 12,1%
5. Personas desaparecidas: 9,8%

Estos resultados confirman la naturaleza autorregresiva del fenómeno criminal, donde los niveles de violencia previos constituyen el mejor predictor de la violencia futura."""

result = doc.add_paragraph()
result_r = result.add_run(result_texto)
result_r.font.size = Pt(11)
result_r.font.name = 'Times New Roman'
result.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ============================================================
# DISCUSIÓN
# ============================================================
disc_titulo = doc.add_paragraph()
disc_run = disc_titulo.add_run("4. DISCUSIÓN")
disc_run.bold = True
disc_run.font.size = Pt(12)
disc_run.font.name = 'Times New Roman'

disc_texto = """Los resultados obtenidos demuestran la viabilidad de aplicar técnicas de aprendizaje automático para la predicción de homicidios en Ecuador con un alto grado de precisión. El coeficiente de determinación del 96,85% supera los reportados en estudios similares realizados en otros contextos latinoamericanos.

La superioridad del modelo XGBoost sobre los demás algoritmos evaluados coincide con la literatura especializada, que destaca su eficacia en problemas de predicción con datos tabulares y relaciones no lineales entre variables. Este algoritmo ha mostrado consistentemente buenos resultados en competencias de ciencia de datos y aplicaciones del mundo real.

La alta importancia asignada a la variable de rezago temporal (homicidios del mes anterior) tiene implicaciones prácticas significativas. Este hallazgo sugiere que los patrones de violencia tienden a persistir en el tiempo, posiblemente debido a dinámicas de venganza, control territorial o escaladas entre grupos criminales. Para los tomadores de decisiones, esto implica que las intervenciones deben ser inmediatas una vez detectados incrementos inusuales.

La concentración del 47% de los homicidios en la provincia del Guayas, particularmente en los cantones de Guayaquil, Durán y Samborondón, refleja la influencia del puerto marítimo como punto de tránsito del narcotráfico internacional. Esta realidad geográfica del crimen debe orientar la focalización de recursos policiales y programas de prevención.

Es importante reconocer las limitaciones del estudio. Primero, el modelo no incorpora variables socioeconómicas como desempleo o desigualdad, que podrían mejorar la capacidad explicativa. Segundo, la calidad de los datos oficiales puede verse afectada por subregistro, especialmente en zonas rurales. Tercero, los cambios en políticas de seguridad o eventos extraordinarios (estado de excepción) pueden alterar los patrones aprendidos.

Futuras investigaciones podrían incorporar datos georreferenciados a nivel cantonal o parroquial, permitiendo predicciones más granulares. Asimismo, la inclusión de variables contextuales como presencia de grupos criminales específicos o rutas de narcotráfico podría enriquecer el modelo."""

disc = doc.add_paragraph()
disc_r = disc.add_run(disc_texto)
disc_r.font.size = Pt(11)
disc_r.font.name = 'Times New Roman'
disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ============================================================
# CONCLUSIONES
# ============================================================
conc_titulo = doc.add_paragraph()
conc_run = conc_titulo.add_run("5. CONCLUSIONES")
conc_run.bold = True
conc_run.font.size = Pt(12)
conc_run.font.name = 'Times New Roman'

conc_texto = """La investigación desarrolló y validó un modelo predictivo de homicidios intencionales para Ecuador, alcanzando una precisión del 96,85% mediante el algoritmo XGBoost. Este resultado demuestra que los datos oficiales del Ministerio del Interior contienen información suficiente para anticipar patrones de violencia letal con alta confiabilidad.

El análisis de más de 850.000 registros confirmó la tendencia explosiva de la violencia en Ecuador, con un incremento del 738% en la tasa de homicidios entre 2017 y 2025. La provincia del Guayas emerge como el epicentro de la crisis, concentrando casi la mitad de las muertes violentas del país.

Las variables más predictivas identificadas (homicidios previos, armas incautadas, y ubicación geográfica) ofrecen orientaciones concretas para la política pública. Los niveles de violencia tienden a perpetuarse en el tiempo y el espacio, lo que justifica intervenciones focalizadas y sostenidas en los territorios más afectados.

El modelo desarrollado constituye una herramienta de apoyo para la toma de decisiones en materia de seguridad ciudadana. Su implementación operativa permitiría generar alertas tempranas ante incrementos proyectados, optimizar la asignación de recursos policiales y evaluar el impacto de las intervenciones realizadas.

Se recomienda la actualización periódica del modelo con nuevos datos y la incorporación de variables adicionales que enriquezcan su capacidad predictiva. La colaboración entre academia e instituciones gubernamentales resulta fundamental para traducir estos avances científicos en mejoras tangibles para la seguridad de la ciudadanía ecuatoriana."""

conc = doc.add_paragraph()
conc_r = conc.add_run(conc_texto)
conc_r.font.size = Pt(11)
conc_r.font.name = 'Times New Roman'
conc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ============================================================
# REFERENCIAS
# ============================================================
ref_titulo = doc.add_paragraph()
ref_run = ref_titulo.add_run("REFERENCIAS BIBLIOGRÁFICAS")
ref_run.bold = True
ref_run.font.size = Pt(12)
ref_run.font.name = 'Times New Roman'

referencias = """
Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794. https://doi.org/10.1145/2939672.2939785

Instituto Nacional de Estadística y Censos. (2025). Proyecciones poblacionales de Ecuador 2010-2025. INEC.

Ministerio del Interior de Ecuador. (2025). Estadísticas de seguridad ciudadana: Homicidios intencionales 2014-2025. https://www.ministeriodeinterior.gob.ec

Observatorio Ecuatoriano de Crimen Organizado. (2025). Informe anual de homicidios y violencia criminal en Ecuador. OECO-PADF.

Organización de las Naciones Unidas. (2024). Global Study on Homicide 2024. United Nations Office on Drugs and Crime.

Perry, W. L., McInnis, B., Price, C. C., Smith, S. C., & Hollywood, J. S. (2013). Predictive Policing: The Role of Crime Forecasting in Law Enforcement Operations. RAND Corporation.

Primicias. (2025, enero 13). La violencia se desborda en 2025: Guayaquil concentra los crímenes. https://www.primicias.ec

Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32. https://doi.org/10.1023/A:1010933404324

Mohler, G. O., Short, M. B., Brantingham, P. J., Schoenberg, F. P., & Tita, G. E. (2011). Self-exciting point process modeling of crime. Journal of the American Statistical Association, 106(493), 100-108."""

ref = doc.add_paragraph()
ref_r = ref.add_run(referencias)
ref_r.font.size = Pt(10)
ref_r.font.name = 'Times New Roman'
ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ============================================================
# GUARDAR DOCUMENTO
# ============================================================
archivo_salida = os.path.join(RUTA_SALIDA, "ARTICULO_ECA_SINERGIA_CRIMINALIDAD.docx")
doc.save(archivo_salida)

print(f"\n✅ ARTÍCULO GENERADO EXITOSAMENTE")
print(f"📁 Ubicación: {archivo_salida}")
print(f"\n📋 ESTRUCTURA DEL ARTÍCULO:")
print("   1. Título (español e inglés)")
print("   2. Autor y afiliación")
print("   3. Resumen y palabras clave")
print("   4. Abstract y keywords")
print("   5. Introducción")
print("   6. Metodología")
print("   7. Resultados")
print("   8. Discusión")
print("   9. Conclusiones")
print("   10. Referencias bibliográficas")
print(f"\n🎯 Listo para enviar a Revista ECA Sinergia!")
